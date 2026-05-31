"""
Genera Memoria_BombChampions_PI.docx desde el borrador Markdown y assets del proyecto.
Ejecutar: .venv\\Scripts\\python.exe PRESENTACION\\ENTREGA\\generar_memoria_docx.py
"""

from __future__ import annotations

import re
import zipfile
import xml.etree.ElementTree as ET
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt

ROOT = Path(__file__).resolve().parents[2]
ENTREGA = Path(__file__).resolve().parent
BORRADOR = ENTREGA / "borradores" / "Memoria_BombChampions_PI.md"
IMAGENES = ENTREGA / "imagenes"
CAPTURAS = IMAGENES / "capturas"
FINAL = ENTREGA / "final"
OUTPUT = FINAL / "Memoria_BombChampions_PI.docx"

EVOLUCION_DOCX = ROOT / "PRESENTACION" / "Capturas de pantallad de evolucion..docx"
FONDO_MENU = ROOT / "assets" / "FONDO BOMB CHAMPIOS.png"
SPRITE_EJEMPLO = ROOT / "assets" / "BLUE ABAJO.png"
TILES_EJEMPLO = ROOT / "assets" / "Pasto.png"
BOMBA_IMG = ROOT / "assets" / "Bomba.png"
EXPLOSION_IMG = ROOT / "assets" / "EXPLOCION.png"


def extract_docx_images(docx_path: Path, out_dir: Path, prefix: str) -> list[Path]:
    out_dir.mkdir(parents=True, exist_ok=True)
    with zipfile.ZipFile(docx_path) as z:
        rels = {}
        for rel in ET.fromstring(z.read("word/_rels/document.xml.rels")):
            target = rel.get("Target")
            if target and "media/" in target:
                rels[rel.get("Id")] = "word/" + target.replace("../", "")
        xml = z.read("word/document.xml").decode("utf-8")
        order = re.findall(r'r:embed="(rId\d+)"', xml)
        saved = []
        for i, rid in enumerate(order, 1):
            media = rels.get(rid)
            if not media:
                continue
            ext = Path(media).suffix
            dest = out_dir / f"{prefix}_{i:02d}{ext}"
            dest.write_bytes(z.read(media))
            saved.append(dest)
    return saved


def add_page_numbers(doc: Document):
    """Numeración centrada en el pie (excepto portada: se aplica a toda la sección)."""
    section = doc.sections[0]
    footer = section.footer
    footer.is_linked_to_previous = False
    p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    fld_char_begin = OxmlElement("w:fldChar")
    fld_char_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    fld_char_sep = OxmlElement("w:fldChar")
    fld_char_sep.set(qn("w:fldCharType"), "separate")
    fld_char_end = OxmlElement("w:fldChar")
    fld_char_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char_begin)
    run._r.append(instr)
    run._r.append(fld_char_sep)
    run._r.append(fld_char_end)
    run.font.name = "Arial"
    run.font.size = Pt(10)


def setup_document() -> Document:
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)
    style = doc.styles["Normal"]
    style.font.name = "Arial"
    style.font.size = Pt(12)
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
    pf = style.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    pf.line_spacing = 1.5
    pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    return doc


def add_para(doc: Document, text: str, *, bold: bool = False, center: bool = False, size: int = 12):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.font.name = "Arial"
    run.font.size = Pt(size)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER if center else WD_ALIGN_PARAGRAPH.JUSTIFY
    return p


def add_heading(doc: Document, text: str, level: int):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.name = "Arial"
        run.font.color.rgb = None
    h.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    h.paragraph_format.line_spacing = 1.5
    return h


def add_bullet(doc: Document, text: str):
    p = doc.add_paragraph(text, style="List Bullet")
    for run in p.runs:
        run.font.name = "Arial"
        run.font.size = Pt(12)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY


def add_numbered(doc: Document, text: str):
    p = doc.add_paragraph(text, style="List Number")
    for run in p.runs:
        run.font.name = "Arial"
        run.font.size = Pt(12)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    p.paragraph_format.line_spacing = 1.5


def add_table(doc: Document, headers: list[str], rows: list[list[str]]):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        for p in hdr[i].paragraphs:
            for run in p.runs:
                run.bold = True
                run.font.name = "Arial"
                run.font.size = Pt(11)
    for r_idx, row in enumerate(rows):
        cells = table.rows[r_idx + 1].cells
        for c_idx, val in enumerate(row):
            cells[c_idx].text = val
            for p in cells[c_idx].paragraphs:
                for run in p.runs:
                    run.font.name = "Arial"
                    run.font.size = Pt(11)
    doc.add_paragraph()


def add_image(doc: Document, path: Path, caption: str, width_cm: float = 14.0):
    if not path.exists():
        add_para(doc, f"[Imagen no encontrada: {path.name}]", center=True)
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(str(path), width=Cm(width_cm))
    cap = doc.add_paragraph(caption)
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in cap.runs:
        run.font.name = "Arial"
        run.font.size = Pt(10)
        run.italic = True
    doc.add_paragraph()


def add_code_block(doc: Document, text: str):
    for line in text.strip().splitlines():
        p = doc.add_paragraph(line)
        p.paragraph_format.left_indent = Cm(0.5)
        for run in p.runs:
            run.font.name = "Courier New"
            run.font.size = Pt(10)


def build_cover(doc: Document):
    doc.add_paragraph()
    if FONDO_MENU.exists():
        add_image(doc, FONDO_MENU, "Figura 1. Fondo del menú principal (Bomb Champions)", width_cm=12)
    add_para(doc, "IES El Álamo", center=True, size=14)
    add_para(doc, "[Localidad — confirmar con el centro]", center=True, size=12)
    doc.add_paragraph()
    add_para(doc, "Proyecto Intermodular SMR", center=True, bold=True, size=13)
    add_para(doc, "(Sistemas Microinformáticos y Redes)", center=True, size=12)
    doc.add_paragraph()
    add_para(doc, "BOMB CHAMPIONS", center=True, bold=True, size=22)
    doc.add_paragraph()
    add_para(doc, "Miguel Eduardo Marcano Ordaz", center=True)
    add_para(doc, "Gabriel Alejandro Celis Ordaz", center=True)
    doc.add_paragraph()
    add_para(doc, "Tutor: Raúl Bernabé Sebastián", center=True)
    doc.add_paragraph()
    add_para(doc, "Fecha de exposición: 5 de junio de 2026", center=True)
    doc.add_page_break()


def build_index(doc: Document):
    add_heading(doc, "2. Índice", 1)
    add_para(
        doc,
        "El índice definitivo con números de página se completará al final, "
        "después de maquetar el documento en Word (Referencias → Tabla de contenido).",
    )
    rows = [
        ("1. Portada", "—"),
        ("2. Índice", "—"),
        ("3. Descripción general del proyecto", "—"),
        ("4. Objetivos", "—"),
        ("5. Justificación de la opción elegida", "—"),
        ("5.1. Uso de inteligencia artificial en el desarrollo", "—"),
        ("6. Recursos necesarios", "—"),
        ("7. Desarrollo del proyecto", "—"),
        ("7.1. Metodología y fases", "—"),
        ("7.2. Arquitectura del software", "—"),
        ("7.3. Desarrollo técnico (Miguel Marcano)", "—"),
        ("7.4. Diseño gráfico y assets (Gabriel Celis)", "—"),
        ("7.5. Música y sonido (Gabriel Celis)", "—"),
        ("7.6. Diseño visual de personajes (Gabriel Celis)", "—"),
        ("7.7. Evolución del juego (capturas)", "—"),
        ("8. Manual básico de utilización", "—"),
        ("9. Dificultades encontradas", "—"),
        ("10. Conclusiones", "—"),
        ("11. Posibles mejoras", "—"),
        ("12. Fuentes de información y webgrafía", "—"),
    ]
    add_table(doc, ["Apartado", "Página"], rows)
    doc.add_page_break()


def build_body(doc: Document, evolucion_imgs: list[Path]):
    add_heading(doc, "3. Descripción general del proyecto", 1)
    for t in [
        "Bomb Champions es un videojuego para ordenador inspirado en el clásico Bomberman. "
        "Dos o más jugadores se mueven por un laberinto, colocan bombas para romper bloques "
        "e intentan eliminar a los demás. El último jugador en pie gana la partida.",
        "La idea principal del proyecto no es solo copiar un juego antiguo, sino añadir "
        "personajes distintos, los campeones, que el jugador elige antes de empezar. Cada "
        "campeón puede tener velocidad, número de bombas o una habilidad especial diferente. "
        "En el Bomberman original todos los personajes son iguales, aquí la elección importa.",
        "El juego se ha programado en Python con la librería de Pygame, de modo que funciona "
        "en cualquier ordenador con tal de instalar python. También incluye partidas en red "
        "local (LAN) entre varios ordenadores, incluyendo asi el módulo de Redes.",
        "El producto está en fase de MVP, que significa producto mínimo viable (Minimum "
        "Viable Product): el menú, el mapa, los jugadores, las bombas y el modo multijugador "
        "por red están implementados. Al desarrollar el juego, se pensó no solo en diseñar lo "
        "mínimo para que sea funcional, si no también dejar el terreno preparado para "
        "escalarlo, o sea, añadir mas mejoras a futuro.",
    ]:
        add_para(doc, t)

    add_heading(doc, "4. Objetivos", 1)
    add_heading(doc, "4.1. Objetivos generales", 2)
    for b in [
        "Desarrollar un videojuego funcional Pygame, aplicando programación orientada a objetos.",
        "Integrar conocimientos de redes (cliente-servidor, sockets) para partidas multijugador en LAN.",
        "Diseñar los assets: Imagenes, sprites, musica y arte del juego.",
    ]:
        add_bullet(doc, b)

    add_heading(doc, "4.2. Objetivos específicos", 2)
    add_table(
        doc,
        ["Objetivo", "Estado"],
        [
            ["Estructura modular del código (main, mapa, bomba, campeones, red_*)", "Logrado"],
            ["Menú principal y flujo de partida local 2 jugadores", "Logrado"],
            ["Generación y dibujo del mapa (17×21 celdas)", "Logrado"],
            ["Sistema de bombas y explosiones", "Logrado"],
            ["Campeones con parámetros y habilidades especiales", "Logrado (en ampliación)"],
            ["Multijugador LAN (crear/buscar partida, sincronización)", "Implementado"],
            ["Assets gráficos y música coherentes con el estilo del juego", "Logrado (Gabriel Celis)"],
            ["Documentar decisiones técnicas y manual de usuario", "En curso (esta memoria)"],
        ],
    )

    add_heading(doc, "5. Justificación de la opción elegida", 1)
    add_para(doc, "Se ha elegido trabajar con Python, Pygame, PyCharm, Piskel y BeepBox por las razones siguientes:")
    for i, t in enumerate(
        [
            "El ciclo SMR incluye programación en Python y fundamentos de redes; el proyecto demuestra ambos.",
            "Python, Pygame, PyCharm Community, Piskel y BeepBox no requieren licencias de pago, se pueden usar gratuitamente.",
            "Piskel encaja con el estilo retro del juego, es facil crear los sprites pixelados que usamos en el proyecto.",
            "BeepBox es una herramienta en linea que facilita mucho la creacion de musica para el juego. Se enfoca en un estilo retro.",
            "Multijugador LAN: Implementacion basica de redes. Se necesita conicimiento de IPs, Firewall, sockets, puertos, y demas para que funcione el juego.",
        ],
        1,
    ):
        add_numbered(doc, t)

    add_heading(doc, "5.1. Uso de inteligencia artificial en el desarrollo", 2)
    add_para(
        doc,
        "El uso de herramientas de inteligencia artificial (asistentes de código, asistencia "
        "para solucionar errores y bugs, ayuda con el diseño de musica.) forma parte "
        "deliberada del proyecto y se justifica en dos ámbitos: formativo y profesional.",
    )
    add_heading(doc, "Justificación formativa", 3)
    for t in [
        "En el instituto ya trabajamos con Python, pero Pygame y las redes con sockets eran cosas nuevas para el equipo. La IA nos ha servido como apoyo cuando nos atascábamos: por ejemplo, entender un error de Pygame, ver cómo organizar el envío de datos por TCP o probar una idea antes de meterla en el código definitivo.",
        "Eso no significa copiar y pegar sin más. Siempre revisamos lo que sale, lo probamos en el juego y lo adaptamos a nuestros archivos (main.py, red_partida.py, etc.). Así hemos aprendido más rápido que leyendo solo documentación, porque veíamos el problema concreto de nuestro proyecto y no un ejemplo genérico.",
        "También nos ha ayudado a redactar borradores de esta memoria y a ordenar ideas para la exposición, pero el contenido final lo escribimos y comprobamos nosotros, sobre todo lo que describe lo que el juego hace de verdad.",
    ]:
        add_para(doc, t)
    add_heading(doc, "Contexto del mercado laboral", 3)
    add_para(
        doc,
        "Las ofertas de empleo para desarrollador de software, técnico de sistemas y perfiles "
        "junior en informática mencionan con frecuencia competencias como:",
    )
    for b in [
        "Productividad con asistentes de IA en el IDE.",
        "Capacidad de revisar y validar código generado o sugerido.",
        "Documentación y pruebas apoyadas en herramientas actuales.",
    ]:
        add_bullet(doc, b)
    add_para(
        doc,
        "Un programador que rechaza por completo el uso de IA en el flujo de trabajo puede resultar "
        "menos competitivo frente a quien la utiliza con criterio: no para sustituir el pensamiento, "
        "sino para acelerar tareas repetitivas, explorar APIs desconocidas (PyGame) y ayuda en la "
        "solucion de errores.",
    )
    add_para(
        doc,
        "En el sector la IA no sustituye al desarrollador: amplía su alcance. La habilidad valorada "
        "es saber qué pedir, comprobar el resultado y asumir la responsabilidad del código y la "
        "documentación entregados.",
    )
    add_heading(doc, "Uso de IA en el trabajo de Gabriel Celis", 3)
    add_para(
        doc,
        "Gabriel ha usado IA solo en una parte concreta: el fondo del menú principal. Para esa "
        "imagen pidió ayuda a ChatGPT y luego la adaptó al estilo del juego. Los personajes, el "
        "mapa, la bomba, las explosiones y el resto de sprites los ha hecho él en Piskel, "
        "cambiando paletas sobre una plantilla encontrada en internet, no generada por IA.",
    )

    add_heading(doc, "6. Recursos necesarios", 1)
    add_heading(doc, "6.1. Hardware", 2)
    add_bullet(doc, "Dos ordenadores portátiles o de sobremesa (uno por integrante del equipo).")
    add_bullet(doc, "Red local Wi‑Fi o cableada para el multijugador LAN.")
    add_heading(doc, "6.2. Software", 2)
    add_table(
        doc,
        ["Herramienta", "Uso"],
        [
            ["Python", "Lenguaje principal"],
            ["Pygame", "Ventana, sprites, eventos, sonido básico"],
            ["PyCharm", "IDE: edición, depuración, entorno virtual"],
            ["Git", "Control de versiones del repositorio del proyecto"],
            ["Piskel", "Sprites pixel art"],
            ["BeepBox", "Música y efectos en estilo retro"],
        ],
    )
    add_heading(doc, "6.3. Servicios en línea", 2)
    add_bullet(doc, "Repositorio del código en Github.")

    add_heading(doc, "7. Desarrollo del proyecto", 1)
    add_heading(doc, "7.1. Metodología y fases", 2)
    fases = [
        "Planificación: Idea del juego, mecánicas, reparto Miguel / Gabriel.",
        "Diseño de estructura: Módulos Python, que sea facil de ver que elementos controla cada archivo, constantes en configuracion.py.",
        "Programación principal: Mapa, jugadores, bombas, menús.",
        "Multijugador: Descubrimiento UDP y partida TCP (red_descubrimiento.py, red_partida.py).",
        "Gráficos y sonido: Assets en assets/ (responsabilidad principal de Gabriel).",
        "Integración y pruebas: Partidas locales y en LAN, corrección de bugs.",
        "Documentación y presentación: Realizar la Memoria (este documento) y presentacion en Canva.",
    ]
    for i, f in enumerate(fases, 1):
        add_numbered(doc, f)

    add_heading(doc, "7.2. Arquitectura del software", 2)
    add_code_block(
        doc,
        """main.py
  ├── configuracion.py    (constantes, teclas, red)
  ├── mapa.py             (clase Mapa: grilla y dibujo)
  ├── bomba.py            (Bomba, GestorBombas)
  ├── campeones.py        (definición de campeones, clase Jugador)
  ├── especiales.py       (habilidades, p. ej. cuchillo)
  ├── red_descubrimiento.py  (anuncio/búsqueda UDP)
  └── red_partida.py         (host/cliente TCP)""",
    )
    add_para(doc, "Estructura de carpetas del repositorio:")
    add_code_block(
        doc,
        """BombChampions/
  main.py, configuracion.py, mapa.py, bomba.py
  campeones.py, especiales.py, red_descubrimiento.py, red_partida.py
  assets/          # imágenes y .piskel
  ejecutar.bat     # lanzador Windows
  requirements.txt""",
    )

    add_heading(doc, "7.3. Desarrollo técnico (Miguel Marcano)", 2)
    add_heading(doc, "Implementado", 3)
    add_table(
        doc,
        ["Componente", "Descripción breve"],
        [
            ["Menú principal", "Título, botones: partida local, crear partida, buscar partidas, salir."],
            ["Mapa", "Grilla 17×21; suelo (0), pared fija (1), caja destructible (2)."],
            ["Jugador", "Movimiento por cuadrícula, colisiones, elección de campeón, sprites."],
            ["Bombas", "Colocación, temporizador, explosión en cruz, destrucción de cajas."],
            ["Campeones", "Azul, Rojo, Amarillo y CuchillasPJ (especial cuchillo)."],
            ["Partida local", "Dos jugadores en el mismo PC (flechas + WASD)."],
            ["Red LAN", "Host anuncia partida; clientes se unen; sincronización desde el host."],
        ],
    )
    add_heading(doc, "Mapa — tipos de celda", 3)
    add_table(
        doc,
        ["Valor", "Significado", "Asset"],
        [
            ["0", "Suelo transitable", "Pasto.png"],
            ["1", "Pared indestructible", "ParedHierro.gif"],
            ["2", "Caja destructible", "ParedLadrillos.png"],
        ],
    )
    add_heading(doc, "Controles (partida local)", 3)
    add_table(
        doc,
        ["Jugador", "Movimiento", "Bomba", "Especial"],
        [
            ["1", "Flechas", "Espacio", "Mayús izquierdo"],
            ["2", "WASD", "E", "Q"],
        ],
    )

    add_heading(doc, "7.4. Diseño gráfico y assets (Gabriel Celis)", 2)
    for t in [
        "La herramienta principal ha sido Piskel. Por experiencia es la más fácil de entender.",
        "Se creo cuatro personajes donde lo que más destaca es la vestimenta: para cada uno ha usado una paleta distinta con tonos claros y oscuros de azul, rojo, amarillo y negro. El boceto base no se hizo la IA, se encontró una plantilla en internet y lo que tocaba era cambiar la paleta y hacer que la animación funcionara en Piskel. El tamaño del sprite es 33×45 píxeles para que el personaje entre entero en el lienzo.",
        "También se encargó del suelo, las paredes de ladrillo (destructibles), los bloques de hierro (indestructibles), la bomba, las explosiones, los corazones y algún item especial.",
    ]:
        add_para(doc, t)
    if SPRITE_EJEMPLO.exists():
        add_image(doc, SPRITE_EJEMPLO, "Figura 2. Ejemplo de sprite de campeón (Guerrero Azul)", width_cm=3)
    row_imgs = [p for p in [TILES_EJEMPLO, BOMBA_IMG, EXPLOSION_IMG] if p.exists()]
    if row_imgs:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for img in row_imgs:
            run = p.add_run()
            run.add_picture(str(img), width=Cm(2.5))
        cap = doc.add_paragraph("Figura 3. Tiles del mapa, bomba y explosión")
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in cap.runs:
            run.font.name = "Arial"
            run.font.size = Pt(10)
            run.italic = True

    add_heading(doc, "7.5. Música y sonido (Gabriel Celis)", 2)
    for t in [
        "Para el audio se ha usado BeepBox. Permite crear instrumentales de forma muy sencilla, en menos de cinco minutos se puede tener algo usable para un juego retro.",
        "Se ha creado los sonidos que el proyecto necesita: instrumental de partida, sonido de la bomba y sonido de la interfaz (menú y botones).",
        "BeepBox es fácil de usar: cada fila de instrumentos lleva un número. Todas las filas marcadas con el mismo número repiten el mismo patrón, si cambias a otro número puedes añadir otra capa sin cortar la primera.",
    ]:
        add_para(doc, t)

    add_heading(doc, "7.6. Diseño visual de personajes / campeones (Gabriel Celis)", 2)
    add_table(
        doc,
        ["Campeón", "Prefijo sprite", "Detalle visual"],
        [
            ["Guerrero Azul", "BLUE", "Paleta azul claro / oscuro"],
            ["Guerrero Rojo", "RED", "Paleta roja"],
            ["Guerrero Amarillo", "YELLOW", "Paleta amarilla"],
            ["CuchillasPJ", "CuchillasPJ", "Chaqueta negra; habilidad especial"],
        ],
    )
    add_para(
        doc,
        "La idea del personaje de chaqueta negra y sus habilidades viene de Jett, de Valorant. "
        "No es una copia: en Valorant Jett usa kunais; en Bomb Champions CuchillasPJ lanza "
        "cuchillos / ondas adaptados al estilo Bomberman.",
    )

    add_heading(doc, "7.7. Evolución del juego (capturas)", 2)
    add_para(
        doc,
        "El juego no salió terminado de golpe. Estas capturas muestran cómo fue mejorando poco a poco.",
    )
    captions = [
        "Figura 4. Versión 1 — MVP jugable (sprites estáticos, interfaz básica)",
        "Figura 5. Versión 1 — partida en curso",
        "Figura 6. Versión 2 — menú con fondo y selección de campeones",
        "Figura 7. Versión 2 — HUD en partida",
        "Figura 8. Versión actual — partida con bombas y explosiones",
        "Figura 9. Versión actual — multijugador / selección",
        "Figura 10. Repositorio del proyecto en GitHub",
    ]
    for i, img in enumerate(evolucion_imgs):
        cap = captions[i] if i < len(captions) else f"Figura {4 + i}. Captura de evolución"
        add_image(doc, img, cap, width_cm=13)

    add_heading(doc, "8. Manual básico de utilización", 1)
    add_heading(doc, "8.1. Requisitos", 2)
    add_bullet(doc, "Ordenador con cualquier sistema operativo que corra Python (Windows, Linux, Mac).")
    add_bullet(doc, "Python 3.10 o superior.")
    add_bullet(doc, "Dependencias en requirements.txt (principalmente pygame).")
    add_heading(doc, "8.2. Instalación", 2)
    add_para(doc, "No requiere instalacion, el juego final es un .exe standalone portable.")
    add_heading(doc, "8.3. Ejecutar el juego", 2)
    add_para(doc, "Doble click en el juego")
    add_heading(doc, "8.4. Menú y modos de juego", 2)
    for t in [
        "INICIAR PARTIDA (local): dos jugadores en el mismo ordenador; elegir campeón y jugar.",
        "CREAR PARTIDA: el host crea una sala LAN; muestra su IP; espera jugadores; INICIAR PARTIDA con mínimo 2.",
        "BUSCAR PARTIDAS: el cliente lista salas en la red y se une a una.",
        "SALIR o tecla ESC para volver al menú o cerrar.",
    ]:
        add_numbered(doc, t)
    add_heading(doc, "8.5. Multijugador LAN — incidencias frecuentes", 2)
    for b in [
        "Ambos PCs en la misma red Wi‑Fi.",
        "Permitir Python en el firewall de Windows (red privada).",
        "Si no aparecen salas, comprobar IP del host y que el puerto no esté bloqueado.",
    ]:
        add_bullet(doc, b)

    add_heading(doc, "9. Dificultades encontradas", 1)
    dificultades = [
        (
            "Aprender a usar GitHub en pareja",
            "Al principio nos costó coordinar el repositorio. Los dos trabajábamos en la misma rama y al hacer git push aparecían conflictos. Tuvimos que aprender a comunicar qué archivo tocaba cada uno, hacer pull antes de subir y, cuando hacía falta, usar ramas.",
        ),
        (
            "Los GIF animados se veían como imágenes fijas",
            "Algunos sprites se exportaron como GIF pensando que Pygame los animaría solos, pero solo se veía el primer fotograma. La solución es exportar cada frame como PNG y mostrarlos uno a uno en código.",
        ),
        (
            "Controles en partida LAN",
            "Al principio el cliente también necesitaba WASD aunque jugara solo a su personaje. Lo corregimos para que cada máquina envíe solo su entrada con las teclas del jugador 1.",
        ),
        (
            "Movimiento encajado al grid",
            "El movimiento parecía teletransporte entre casillas. Separamos posición lógica en el grid de posición en píxeles, interpolando entre celdas.",
        ),
        (
            "Diseño gráfico y audio",
            "Pasar de la versión 1 sin interfaz a menú con fondo, selección de campeones y HUD sin desordenar el estilo pixel art. BeepBox ayudó con el sonido; integrarlo en Pygame sigue en curso.",
        ),
    ]
    for title, body in dificultades:
        add_heading(doc, title, 2)
        add_para(doc, body)

    add_heading(doc, "10. Conclusiones", 1)
    add_heading(doc, "Miguel Marcano (programación y arquitectura)", 2)
    add_para(
        doc,
        "Hemos aprendido cosas importantes sobre el desarollo de juegos y aplicaciones. "
        "Respecto al codigo, a como estructurarlo de forma modular, osea, distintos archivos "
        "python que cumplan roles distintos, y que se comuniquen entre ellos. Se aprendio a "
        "implementar librerias desconocidas (Pygame) en el proyecto, programacion orientada a "
        "objetos, uso de listas y implementacion de redes en el codigo.",
    )
    add_heading(doc, "Gabriel Celis (diseño gráfico, audio y personajes)", 2)
    add_para(
        doc,
        "He aprendido a sacarle partido a Piskel para sprites y animaciones, y a BeepBox para "
        "montar música y efectos en poco tiempo. Lo más importante ha sido coordinar con Miguel: "
        "mis PNG y sonidos tienen que encajar con lo que el código espera en assets/. El fondo "
        "del menú lo hice con ayuda de ChatGPT; el resto del arte del juego lo he hecho yo a "
        "mano en Piskel.",
    )
    add_heading(doc, "Conclusión conjunta", 2)
    add_para(
        doc,
        "Bomb Champions une programación, redes y diseño en un mismo proyecto. No está cerrado "
        "como un juego comercial, pero ya se puede jugar en local y por LAN, y queda una base "
        "clara para seguir con más campeones, sonido y mejoras después del PI.",
    )

    add_heading(doc, "11. Posibles mejoras", 1)
    for b in [
        "Nuevos campeones con habilidades distintas (escudo, teletransporte, etc.).",
        "IA para enemigos en modo solitario.",
        "Mapas alternativos y editor de niveles.",
        "Interfaz más pulida (animaciones de menú, tutorial integrado).",
    ]:
        add_bullet(doc, b)

    add_heading(doc, "12. Fuentes de información y webgrafía", 1)
    add_table(
        doc,
        ["Referencia", "URL"],
        [
            ["BeepBox", "https://www.beepbox.co/"],
            ["Documentación Pygame", "https://www.pygame.org/docs/"],
            ["JetBrains — PyCharm", "https://www.jetbrains.com/pycharm/"],
            ["Piskel — editor de sprites", "https://www.piskelapp.com/"],
            ["Python — documentación oficial", "https://docs.python.org/3/"],
            ["Wikipedia — Bomberman", "https://es.wikipedia.org/wiki/Bomberman"],
            ["Wikipedia — Valorant", "https://es.wikipedia.org/wiki/Valorant"],
        ],
    )


def main():
    CAPTURAS.mkdir(parents=True, exist_ok=True)
    FINAL.mkdir(parents=True, exist_ok=True)

    evolucion_imgs: list[Path] = []
    if EVOLUCION_DOCX.exists():
        evolucion_imgs = extract_docx_images(EVOLUCION_DOCX, CAPTURAS, "evolucion")
        print(f"Extraídas {len(evolucion_imgs)} capturas de evolución")

    doc = setup_document()
    build_cover(doc)
    build_index(doc)
    build_body(doc, evolucion_imgs)
    add_page_numbers(doc)
    doc.save(OUTPUT)
    print(f"Generado: {OUTPUT}")
    print(f"Tamaño: {OUTPUT.stat().st_size // 1024} KB")


if __name__ == "__main__":
    main()
